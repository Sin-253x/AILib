from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

# ======================== 代码解释 ========================
# 1. 整体功能：
#    解析上传的文本、PDF 和 DOCX 文档，把文件字节转换成可保存的文档内容和来源元数据。
#
# 2. 关键部分拆解：
#    - ParsedDocument：保存解析后的标题、正文和来源信息。
#    - parse_uploaded_document：执行扩展名、大小、编码和空内容校验。
#    - _decode_text：兼容 UTF-8 和 UTF-8 BOM 文本。
#    - _extract_pdf_text/_extract_docx_text：通过专用库抽取办公文档正文。
#
# 3. 重要概念与库：
#    - dataclass：用轻量结构承载解析结果。
#    - pathlib.Path：安全提取文件扩展名和无后缀标题。
#    - pypdf/python-docx：分别负责 PDF 和 Word 文档的正文读取。
#
# 4. 潜在问题与改进建议：
#    - 扫描版 PDF 可能没有可抽取文本；需要 OCR 时应另接异步任务。
#    - 当前整篇保存正文；向量化阶段需要继续切分为 chunks。
#
# 5. 修改指南：
#    - 如果要支持新扩展名，建议先更新 SUPPORTED_EXTENSIONS 和对应测试。
# ========================================================
SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}
SUPPORTED_BINARY_EXTENSIONS = {".pdf", ".docx"}
SUPPORTED_EXTENSIONS = SUPPORTED_TEXT_EXTENSIONS | SUPPORTED_BINARY_EXTENSIONS


class DocumentParseError(ValueError):
    """Raised when an uploaded document cannot be accepted or parsed."""


@dataclass(frozen=True)
class ParsedDocument:
    title: str
    content: str
    source_filename: str
    source_mime_type: str | None
    source_size_bytes: int


def _decode_text(content_bytes: bytes) -> str:
    try:
        return content_bytes.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentParseError("File must be valid UTF-8 text") from exc


# ======================== 代码解释 ========================
# 1. 整体功能：
#    从 PDF 字节中抽取可索引文本。
#
# 2. 关键部分拆解：
#    - BytesIO：让 pypdf 像读取文件一样读取内存字节。
#    - PdfReader.pages：逐页抽取文本并拼接。
#
# 3. 重要概念与库：
#    - pypdf：纯 Python PDF 解析库，适合抽取数字文本层。
#
# 4. 潜在问题与改进建议：
#    - 图片扫描 PDF 不含文本层，后续需要 OCR 才能进入知识库。
#
# 5. 修改指南：
#    - 如果要保留页码引用，可在这里返回带页码标记的文本片段。
# ========================================================
def _extract_pdf_text(content_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentParseError("PDF support requires the pypdf package") from exc

    try:
        reader = PdfReader(BytesIO(content_bytes))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise DocumentParseError("Unable to parse PDF file") from exc


# ======================== 代码解释 ========================
# 1. 整体功能：
#    从 DOCX 字节中抽取段落和表格文本。
#
# 2. 关键部分拆解：
#    - Document：python-docx 的 Word 文档读取入口。
#    - paragraphs/tables：覆盖正文段落和常见表格内容。
#
# 3. 重要概念与库：
#    - python-docx：读取 Office Open XML Word 文档的常用库。
#
# 4. 潜在问题与改进建议：
#    - 当前不解析页眉页脚、批注和嵌入对象；后续可按需要补充。
#
# 5. 修改指南：
#    - 如果要保留表格结构，可把行列分隔符调整为更适合 RAG 的 Markdown 表格。
# ========================================================
def _extract_docx_text(content_bytes: bytes) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise DocumentParseError("DOCX support requires the python-docx package") from exc

    try:
        document = DocxDocument(BytesIO(content_bytes))
    except Exception as exc:
        raise DocumentParseError("Unable to parse DOCX file") from exc

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def parse_uploaded_document(
    *,
    filename: str,
    content_type: str | None,
    content_bytes: bytes,
    max_size_bytes: int,
) -> ParsedDocument:
    source_size_bytes = len(content_bytes)
    if source_size_bytes > max_size_bytes:
        raise DocumentParseError("Uploaded file is too large")

    source_path = Path(filename)
    extension = source_path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError("Unsupported file type")

    if extension == ".pdf":
        content = _extract_pdf_text(content_bytes).strip()
    elif extension == ".docx":
        content = _extract_docx_text(content_bytes).strip()
    else:
        content = _decode_text(content_bytes).strip()
    if not content:
        raise DocumentParseError("Uploaded file is empty")

    title = source_path.stem.strip() or "Untitled document"
    return ParsedDocument(
        title=title[:200],
        content=content,
        source_filename=source_path.name,
        source_mime_type=content_type,
        source_size_bytes=source_size_bytes,
    )
